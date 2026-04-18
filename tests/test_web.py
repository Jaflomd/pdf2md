from __future__ import annotations

import io
import zipfile
from pathlib import Path

import fitz
from docx import Document

from pdf2md.web import convert_document_bytes


def _make_pdf_bytes() -> bytes:
    pdf = fitz.open()
    page_1 = pdf.new_page()
    page_1.insert_text((72, 72), "Chapter 1", fontsize=24)
    page_1.insert_text((72, 120), "This is content for chapter 1.", fontsize=12)

    page_2 = pdf.new_page()
    page_2.insert_text((72, 72), "Chapter 2", fontsize=24)
    page_2.insert_text((72, 120), "This is content for chapter 2.", fontsize=12)

    pdf.set_toc([[1, "Chapter 1", 1], [1, "Chapter 2", 2]])
    return pdf.tobytes()


def _make_docx_bytes(tmp_path: Path) -> bytes:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("This is content for chapter 1.")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("This is content for chapter 2.")
    doc.save(path)
    return path.read_bytes()


def test_convert_document_bytes_returns_markdown_and_bundle(tmp_path: Path) -> None:
    result = convert_document_bytes(
        file_bytes=_make_pdf_bytes(),
        filename="sample.pdf",
        chunk_target=80,
        chunk_overlap=10,
        engine="pymupdf4llm",
    )

    bundle_name = result.manifest["output"]["bundle_name"]

    assert "Chapter 1" in result.document_markdown
    assert result.manifest["source"]["path"] == "sample.pdf"
    assert result.manifest["engine"]["full_markdown_engine"] == "pymupdf4llm"
    assert result.archive_name == f"{bundle_name}.zip"
    assert len(result.chapter_markdowns) == 2

    with zipfile.ZipFile(io.BytesIO(result.archive_bytes)) as archive:
        names = set(archive.namelist())

    assert f"{bundle_name}/document.md" in names
    assert f"{bundle_name}/manifest.json" in names


def test_convert_document_bytes_supports_docx(tmp_path: Path) -> None:
    result = convert_document_bytes(
        file_bytes=_make_docx_bytes(tmp_path),
        filename="sample.docx",
        chunk_target=80,
        chunk_overlap=10,
        engine="pymupdf4llm",
    )

    assert "Chapter 1" in result.document_markdown
    assert result.manifest["source"]["path"] == "sample.docx"
    assert result.manifest["source"]["input_format"] == "docx"
    assert result.manifest["output"]["root_path"] == result.manifest["output"]["bundle_name"]


def test_convert_document_bytes_supports_fast_mode(tmp_path: Path) -> None:
    result = convert_document_bytes(
        file_bytes=_make_pdf_bytes(),
        filename="sample.pdf",
        fast_mode=True,
        page_group_size=1,
    )

    assert result.manifest["processing"]["fast_mode"] is True
    assert result.manifest["processing"]["page_group_size"] == 1
    assert result.manifest["chunking"]["enabled"] is True
    assert len(result.chapter_markdowns) == 2
