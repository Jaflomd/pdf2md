from __future__ import annotations

import json
from pathlib import Path

import fitz
import pytest
from docx import Document

from pdf2md.pipeline import UnsupportedInputError, run_conversion


def _make_pdf(path: Path) -> None:
    pdf = fitz.open()
    page_1 = pdf.new_page()
    page_1.insert_text((72, 72), "Chapter 1", fontsize=24)
    page_1.insert_text((72, 120), "This is content for chapter 1.", fontsize=12)

    page_2 = pdf.new_page()
    page_2.insert_text((72, 72), "Chapter 2", fontsize=24)
    page_2.insert_text((72, 120), "This is content for chapter 2.", fontsize=12)

    pdf.set_toc([[1, "Chapter 1", 1], [1, "Chapter 2", 2]])
    pdf.save(path)


def _make_multi_page_pdf(path: Path, pages: int, with_toc: bool = True, heading_size: int = 24) -> None:
    pdf = fitz.open()
    toc: list[list[int | str]] = []
    for page_index in range(1, pages + 1):
        page = pdf.new_page()
        page.insert_text((72, 72), f"Page {page_index}", fontsize=heading_size)
        page.insert_text((72, 120), f"This is page {page_index}.", fontsize=12)
        if with_toc:
            toc.append([1, f"Page {page_index}", page_index])
    if toc:
        pdf.set_toc(toc)
    pdf.save(path)


def _make_plain_pdf(path: Path, pages: int) -> None:
    pdf = fitz.open()
    for page_index in range(1, pages + 1):
        page = pdf.new_page()
        page.insert_text((72, 72), f"This is plain body text for page {page_index}.", fontsize=12)
        page.insert_text((72, 104), "It should not look like a chapter heading.", fontsize=12)
    pdf.save(path)


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("This is content for chapter 1.")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("This is content for chapter 2.")
    doc.save(path)


def test_run_conversion_writes_expected_outputs_and_manifest(tmp_path: Path) -> None:
    source = tmp_path / "sample.pdf"
    outdir = tmp_path / "out"
    _make_pdf(source)

    result = run_conversion(
        input_path=source,
        outdir=outdir,
        chunk_target=80,
        chunk_overlap=10,
        engine="pymupdf4llm",
    )

    root = result.output_root
    manifest = json.loads((root / "manifest.json").read_text(encoding="utf-8"))

    assert root.name.startswith("sample--")
    assert (root / "document.md").exists()
    assert (root / "chapters" / "01-chapter-1.md").exists()
    assert (root / "chapters" / "02-chapter-2.md").exists()
    assert (root / "chunks" / "index.jsonl").exists()
    assert any(path.name.startswith("chunk-") for path in (root / "chunks" / "01-chapter-1").iterdir())
    assert manifest["engine"]["full_markdown_engine"] == "pymupdf4llm"
    assert manifest["output"]["bundle_name"] == root.name
    assert manifest["chunking"]["enabled"] is True
    assert len(manifest["chapters"]) == 2


def test_run_conversion_uses_deterministic_bundle_name_and_cleans_previous_output(tmp_path: Path) -> None:
    source = tmp_path / "sample.pdf"
    outdir = tmp_path / "out"
    _make_pdf(source)

    first = run_conversion(input_path=source, outdir=outdir, engine="pymupdf4llm")
    stale_path = first.output_root / "stale.txt"
    stale_path.write_text("stale", encoding="utf-8")

    second = run_conversion(input_path=source, outdir=outdir, engine="pymupdf4llm")

    assert first.output_root == second.output_root
    assert not stale_path.exists()


def test_auto_engine_falls_back_to_pymupdf_when_docling_errors(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "fallback.pdf"
    outdir = tmp_path / "out"
    _make_pdf(source)

    from pdf2md import pipeline

    def fail_docling(_: Path) -> tuple[str, list[dict[str, int | str | None]]]:
        raise RuntimeError("docling unavailable in test")

    monkeypatch.setattr(pipeline, "_extract_full_markdown_with_docling", fail_docling)

    result = run_conversion(
        input_path=source,
        outdir=outdir,
        chunk_target=80,
        chunk_overlap=10,
        engine="auto",
    )

    assert result.manifest["engine"]["fallback_used"] is True
    assert result.manifest["engine"]["full_markdown_engine"] == "pymupdf4llm"
    assert result.manifest["warnings"]


def test_auto_split_falls_back_to_page_batches_when_chapters_are_not_detected(tmp_path: Path) -> None:
    source = tmp_path / "plain.pdf"
    outdir = tmp_path / "out"
    _make_plain_pdf(source, pages=4)

    result = run_conversion(
        input_path=source,
        outdir=outdir,
        engine="pymupdf4llm",
        split_mode="auto",
        page_group_size=2,
    )

    assert result.manifest["processing"]["segmentation_mode"] == "page-batch"
    assert len(result.manifest["chapters"]) == 2
    assert any("fell back to 2-page sections" in warning for warning in result.manifest["warnings"])


def test_azw3_requires_calibre_when_converter_missing(tmp_path: Path) -> None:
    source = tmp_path / "book.azw3"
    source.write_bytes(b"not-a-real-kindle-file")

    with pytest.raises(UnsupportedInputError):
        run_conversion(
            input_path=source,
            outdir=tmp_path / "out",
            engine="auto",
        )


def test_run_conversion_with_docx_support_writes_expected_outputs(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    outdir = tmp_path / "out"
    _make_docx(source)

    result = run_conversion(
        input_path=source,
        outdir=outdir,
        chunk_target=80,
        chunk_overlap=10,
        engine="pymupdf4llm",
    )

    root = result.output_root
    manifest = json.loads((root / "manifest.json").read_text(encoding="utf-8"))

    assert (root / "document.md").exists()
    assert (root / "chapters" / "01-chapter-1.md").exists()
    assert (root / "chapters" / "02-chapter-2.md").exists()
    assert manifest["source"]["input_format"] == "docx"
    assert manifest["engine"]["full_markdown_engine"] == "pymupdf4llm"
    assert len(manifest["chapters"]) == 2


def test_docx_accepts_docling_engine(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = tmp_path / "sample.docx"
    outdir = tmp_path / "out"
    _make_docx(source)

    from pdf2md import pipeline

    def fake_docling(_: Path) -> tuple[str, list[dict[str, int | str | None]]]:
        markdown = "# Chapter 1\n\nThis is content for chapter 1.\n\n# Chapter 2\n\nThis is content for chapter 2.\n"
        return markdown, []

    monkeypatch.setattr(pipeline, "_extract_full_markdown_with_docling", fake_docling)

    result = run_conversion(
        input_path=source,
        outdir=outdir,
        chunk_target=80,
        chunk_overlap=10,
        engine="docling",
    )

    assert result.manifest["engine"]["full_markdown_engine"] == "docling"


def test_fast_mode_groups_output_by_page_count_and_still_chunks(tmp_path: Path) -> None:
    source = tmp_path / "fast.pdf"
    outdir = tmp_path / "out"
    _make_multi_page_pdf(source, pages=5)

    result = run_conversion(
        input_path=source,
        outdir=outdir,
        fast_mode=True,
        page_group_size=2,
        engine="docling",
    )

    root = result.output_root
    index_lines = (root / "chunks" / "index.jsonl").read_text(encoding="utf-8").strip().splitlines()
    manifest = json.loads((root / "manifest.json").read_text(encoding="utf-8"))

    assert (root / "chapters" / "01-pages-001-002.md").exists()
    assert (root / "chapters" / "02-pages-003-004.md").exists()
    assert (root / "chapters" / "03-page-005.md").exists()
    assert index_lines
    assert manifest["processing"]["fast_mode"] is True
    assert manifest["processing"]["page_group_size"] == 2
    assert manifest["processing"]["segmentation_mode"] == "page-batch"
    assert manifest["chunking"]["enabled"] is True
    assert manifest["engine"]["full_markdown_engine"] == "pymupdf4llm"
    assert any("Fast mode" in warning for warning in manifest["warnings"])
